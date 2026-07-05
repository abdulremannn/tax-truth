import os
from docx import Document as DocxDocument
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from fpdf import FPDF

GREEN = RGBColor(0x10, 0x8A, 0x00)


def export_docx(report: dict, client_name: str, out_path: str):
    doc = DocxDocument()

    title = doc.add_heading("Tax Strategy Consultancy Report", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_paragraph(f"Prepared for: {client_name}")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading("Executive Summary", level=1)
    doc.add_paragraph(report.get("executive_summary", ""))

    doc.add_heading("Recommended Strategies", level=1)
    for s in report.get("strategies", []):
        doc.add_heading(s.get("name", ""), level=2)
        doc.add_paragraph(f"What it is: {s.get('what_it_is', '')}")
        doc.add_paragraph(f"Why you qualify: {s.get('why_you_qualify', '')}")
        doc.add_paragraph(f"Estimated impact: {s.get('estimated_impact', '')}")
        doc.add_paragraph(f"Implementation notes: {s.get('implementation_notes', '')}")

    doc.add_heading("Prioritized Action Plan", level=1)
    for a in report.get("action_plan", []):
        doc.add_paragraph(f"{a.get('priority')}. {a.get('strategy')} — {a.get('reason')}")

    doc.add_heading("Disclaimer", level=1)
    doc.add_paragraph(report.get("disclaimer", ""))

    doc.save(out_path)
    return out_path


class PDFReport(FPDF):
    def header(self):
        self.set_font("Helvetica", "B", 16)
        self.set_text_color(16, 138, 0)
        self.cell(0, 10, "Tax Strategy Consultancy Report", ln=True, align="C")
        self.ln(2)

    def footer(self):
        self.set_y(-15)
        self.set_font("Helvetica", "I", 8)
        self.set_text_color(120, 120, 120)
        self.cell(0, 10, f"Page {self.page_no()}", align="C")


def _clean(text: str) -> str:
    return (text or "").encode("latin-1", "replace").decode("latin-1")


def export_pdf(report: dict, client_name: str, out_path: str):
    pdf = PDFReport()
    pdf.set_margins(15, 15, 15)
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()

    pdf.set_font("Helvetica", "", 11)
    pdf.set_text_color(80, 80, 80)
    pdf.cell(0, 8, _clean(f"Prepared for: {client_name}"), ln=True, align="C")
    pdf.ln(6)

    pdf.set_font("Helvetica", "B", 13)
    pdf.set_text_color(0, 0, 0)
    pdf.cell(0, 8, "Executive Summary", ln=True)
    pdf.set_font("Helvetica", "", 10)
    pdf.multi_cell(0, 6, _clean(report.get("executive_summary", "")))
    pdf.ln(4)

    pdf.set_font("Helvetica", "B", 13)
    pdf.cell(0, 8, "Recommended Strategies", ln=True)
    for s in report.get("strategies", []):
        pdf.set_x(pdf.l_margin)
        pdf.set_font("Helvetica", "B", 11)
        pdf.set_text_color(16, 106, 0)
        pdf.multi_cell(0, 7, _clean(s.get("name", "")))
        pdf.set_text_color(0, 0, 0)
        pdf.set_font("Helvetica", "", 10)
        pdf.set_x(pdf.l_margin)
        pdf.multi_cell(0, 6, _clean(f"What it is: {s.get('what_it_is', '')}"))
        pdf.set_x(pdf.l_margin)
        pdf.multi_cell(0, 6, _clean(f"Why you qualify: {s.get('why_you_qualify', '')}"))
        pdf.set_x(pdf.l_margin)
        pdf.multi_cell(0, 6, _clean(f"Estimated impact: {s.get('estimated_impact', '')}"))
        pdf.set_x(pdf.l_margin)
        pdf.multi_cell(0, 6, _clean(f"Implementation notes: {s.get('implementation_notes', '')}"))
        pdf.ln(3)

    pdf.set_font("Helvetica", "B", 13)
    pdf.cell(0, 8, "Prioritized Action Plan", ln=True)
    pdf.set_font("Helvetica", "", 10)
    for a in report.get("action_plan", []):
        pdf.set_x(pdf.l_margin)
        pdf.multi_cell(0, 6, _clean(f"{a.get('priority')}. {a.get('strategy')} - {a.get('reason')}"))
    pdf.ln(4)

    pdf.set_font("Helvetica", "I", 9)
    pdf.set_text_color(120, 90, 0)
    pdf.multi_cell(0, 6, _clean(report.get("disclaimer", "")))

    pdf.output(out_path)
    return out_path